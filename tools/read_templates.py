# -*- coding: utf-8 -*-
import os
from docx import Document

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "模板文件")
OUTPUT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates_output.txt")

files = [
    os.path.join(TEMPLATE_DIR, "项目组", "邀请函.docx"),
    os.path.join(TEMPLATE_DIR, "项目组", "答复函.docx"),
    os.path.join(TEMPLATE_DIR, "人事组", "邀请函.docx"),
    os.path.join(TEMPLATE_DIR, "人事组", "答复函.docx"),
    os.path.join(TEMPLATE_DIR, "运营组", "邀请函.docx"),
    os.path.join(TEMPLATE_DIR, "运营组", "答复函.docx"),
]

with open(OUTPUT_FILE, 'w', encoding='utf-8') as out:
    for f in files:
        out.write('=' * 80 + '\n')
        out.write(f'文件: {f}\n')
        out.write('=' * 80 + '\n')
        doc = Document(f)

        out.write('--- 段落内容 ---\n')
        for i, p in enumerate(doc.paragraphs):
            text = p.text
            out.write(f'  段落{i}: {text}\n')

        out.write('\n--- 表格内容 ---\n')
        for ti, table in enumerate(doc.tables):
            out.write(f'  表格{ti}:\n')
            for ri, row in enumerate(table.rows):
                cells = [cell.text.replace('\n', '\\n').replace('\r', '') for cell in row.cells]
                out.write(f'    行{ri}: {cells}\n')

        out.write('\n\n')
