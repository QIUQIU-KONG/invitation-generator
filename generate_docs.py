# -*- coding: utf-8 -*-
"""
从飞书多维表格读取出差数据，查询 SQLite 知识库，生成邀请函和答复函 (.docx)

用法:
    python generate_docs.py                          # 从飞书表格读取
    python generate_docs.py path/to/your_file.xlsx   # 从本地 Excel 读取

依赖:
    pip install openpyxl python-docx pandas python-dateutil
"""
import sys
import os
import json
import hashlib
import subprocess
import sqlite3
from dateutil.relativedelta import relativedelta
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from pinyin_data import PINYIN_DATA

# === 配置 ===
DEFAULT_EXCEL = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sample_data.xlsx")
TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "模板文件")
OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
DB_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "knowledge_base.db")
TRACKING_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "generated.json")
AUTO_YES = False
DRY_RUN = False
FORCE_REGENERATE = False
ALLOW_MISSING_PASSPORT = False

# 飞书多维表格配置
FEISHU_BASE_TOKEN = os.environ.get("FEISHU_BASE_TOKEN", "")
FEISHU_TABLE_ID = os.environ.get("FEISHU_TABLE_ID", "")
FEISHU_ADMIN_ID = os.environ.get("FEISHU_ADMIN_ID", "")
LARK_CLI = os.environ.get("LARK_CLI_PATH", "lark-cli")

DEPT_TEMPLATE_MAP = {
    "项目组": "项目组", "项目": "项目组",
    "人事组": "人事组", "人事": "人事组",
    "运营组": "运营组", "运营": "运营组",
    "供应商": "供应商", "供应": "供应商",
}
TEMPLATE_TYPE_KEYS = ("模板类型", "模板", "template_type", "template", "templateType")

# === 字体配置 ===
FONT_NAME = 'Century Gothic'
FONT_SIZE_PT = 12  # 小四

# === 职位翻译表 ===
POSITION_MAP = {
    "经理": "Manager", "总经理": "General Manager", "副总经理": "Deputy General Manager",
    "总监": "Director", "主管": "Supervisor", "助理": "Assistant",
    "工程师": "Engineer", "高级工程师": "Senior Engineer", "初级工程师": "Junior Engineer",
    "技术员": "Technician", "操作员": "Operator", "分析师": "Analyst",
    "顾问": "Consultant", "专员": "Specialist", "协调员": "Coordinator",
    "主任": "Chief", "组长": "Team Leader", "副组长": "Deputy Team Leader",
    "秘书": "Secretary", "翻译": "Translator", "会计": "Accountant",
    "设计师": "Designer", "项目经理": "Project Manager",
    "质检员": "Quality Inspector", "采购员": "Purchaser", "销售": "Sales",
    "实习生": "Intern", "技师": "Technician", "切割技师": "Cutting Technician",
    "多线切割技师": "Multi-wire Cutting Technician", "中级操作员": "Intermediate Operator",
    "体系工程师": "System Engineer", "副经理": "Deputy Manager",
    "设备工程师": "Equipment Engineer",
    "资深基建工程师": "Senior Infrastructure Engineer",
    "项目副总监": "Deputy Project Director",
    "资深制程巡检员": "Senior Process Inspector",
    "资深测试技术工程师": "Senior Test Technology Engineer",
    "过程质量工程师": "Process Quality Engineer",
    "制造技术副经理": "Deputy Manufacturing Technology Manager",
    "电镀工程师": "Plating Engineer",
    "PME工程师": "PME Engineer",
    "线长": "Line Leader",
}


# === 飞书通知 ===
def notify_feishu(title, message):
    """通过飞书发送通知消息"""
    import subprocess
    try:
        content = f"[{title}]\n{message}"
        result = subprocess.run(
            [LARK_CLI, "im", "+messages-send",
             "--user-id", FEISHU_ADMIN_ID,
             "--text", content,
             "--as", "user"],
            capture_output=True, text=True, timeout=10,
            encoding='utf-8', errors='replace'
        )
        if result.returncode == 0:
            print("  [OK] 飞书通知已发送")
        else:
            print(f"  [WARN] 飞书通知发送失败: {result.stderr[:200]}")
    except Exception as e:
        print(f"  [警告] 飞书通知异常: {e}")


def update_feishu_status(record_id, status):
    """更新飞书表格中的生成状态"""
    if not record_id:
        return
    try:
        result = subprocess.run(
            [LARK_CLI, "base", "+record-upsert",
             "--base-token", FEISHU_BASE_TOKEN,
             "--table-id", FEISHU_TABLE_ID,
             "--record-id", record_id,
             "--json", json.dumps({"生成状态": status}, ensure_ascii=False),
             "--as", "user"],
            capture_output=True, text=True, timeout=10,
            encoding='utf-8', errors='replace'
        )
        if result.returncode == 0:
            print(f"  [OK] {record_id}: 已标记为{status}")
        else:
            print(f"  [WARN] {record_id}: 状态更新失败 - {result.stderr[:200]}")
    except Exception as e:
        print(f"  [警告] 飞书状态更新失败: {e}")


# === 数据读取 ===
def read_feishu_table(base_token, table_id):
    """从飞书多维表格读取出差数据"""
    cmd = [
        LARK_CLI, "base", "+record-list",
        "--base-token", base_token,
        "--table-id", table_id,
        "--format", "json",
        "--as", "user"
    ]
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=30, encoding='utf-8', errors='replace')
        if result.returncode != 0:
            message = (result.stderr or result.stdout or "").strip()
            print("[错误] 飞书读取失败")
            if "need_user_authorization" in message or "base:record:read" in message:
                print("原因：飞书账号还没有授权读取多维表格，请先完成 lark-cli 用户授权。")
            if message:
                print(message[:1000])
            return None
        parsed = parse_feishu_json(result.stdout)
        if parsed is not None:
            return parsed
        print("[警告] 飞书返回内容不是预期 JSON，改用表格文本解析。")
        records = []
        lines = result.stdout.strip().split('\n')
        data_start = False
        for line in lines:
            line = line.strip()
            if line.startswith('|') and not data_start:
                data_start = True
                continue
            if data_start and line.startswith('|') and not line.startswith('| ---'):
                if 'Meta:' in line:
                    continue
                if '记录 ID' in line or 'record_id' in line.lower():
                    continue
                cells = [c.strip() for c in line.split('|')[1:-1]]
                if len(cells) >= 5:
                    records.append({
                        'record_id': cells[0],
                        '工号': cells[1].zfill(3),
                        '护照号': cells[2] if cells[2] else 'N/A',
                        '出发时间': cells[3],
                        '返程时间': cells[4],
                        '模板类型': cells[5] if len(cells) >= 6 else ''
                    })
        return records
    except Exception as e:
        print(f"[错误] 飞书读取异常: {e}")
        return None


def unwrap_feishu_cell(value):
    """Convert common Feishu Base cell shapes into plain text."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, (int, float)):
        return str(int(value)) if isinstance(value, float) and value.is_integer() else str(value)
    if isinstance(value, list):
        return "".join(unwrap_feishu_cell(item) for item in value).strip()
    if isinstance(value, dict):
        for key in ("text", "name", "value", "title", "email", "phone"):
            if key in value:
                return unwrap_feishu_cell(value[key])
        if "timestamp" in value:
            return unwrap_feishu_cell(value["timestamp"])
    return str(value).strip()


def first_value(mapping, keys):
    for key in keys:
        if key in mapping and mapping.get(key) not in (None, ""):
            return mapping.get(key)
    return ""


def parse_feishu_json(stdout):
    try:
        payload = json.loads(stdout)
    except Exception:
        return None

    data = payload.get("data", payload) if isinstance(payload, dict) else payload
    if isinstance(data, dict) and isinstance(data.get("data"), list) and isinstance(data.get("fields"), list):
        fields = data.get("fields") or []
        record_ids = data.get("record_id_list") or []
        records = []
        for idx, row in enumerate(data.get("data") or []):
            if not isinstance(row, list):
                continue
            values = {field: unwrap_feishu_cell(row[pos]) if pos < len(row) else "" for pos, field in enumerate(fields)}
            emp_id = values.get("工号") or values.get("员工号") or values.get("emp_id") or ""
            records.append({
                "record_id": record_ids[idx] if idx < len(record_ids) else "",
                "工号": emp_id.zfill(3) if emp_id else "",
                "护照号": values.get("护照号") or values.get("passport") or "N/A",
                "出发时间": values.get("出发时间") or values.get("出发日期") or values.get("depart_date") or "",
                "返程时间": values.get("返程时间") or values.get("返程日期") or values.get("return_date") or "",
                "模板类型": first_value(values, TEMPLATE_TYPE_KEYS),
                "生成状态": values.get("生成状态") or values.get("status") or "",
            })
        return records

    items = []
    if isinstance(data, dict):
        items = data.get("items") or data.get("records") or data.get("record_list") or []
    elif isinstance(data, list):
        items = data

    records = []
    for item in items:
        if not isinstance(item, dict):
            continue
        fields = item.get("fields") or item.get("record", {}).get("fields") or {}
        if not isinstance(fields, dict):
            continue
        emp_id = unwrap_feishu_cell(fields.get("工号") or fields.get("员工号") or fields.get("emp_id"))
        depart_date = unwrap_feishu_cell(fields.get("出发时间") or fields.get("出发日期") or fields.get("depart_date"))
        return_date = unwrap_feishu_cell(fields.get("返程时间") or fields.get("返程日期") or fields.get("return_date"))
        passport = unwrap_feishu_cell(fields.get("护照号") or fields.get("passport"))
        template_type = unwrap_feishu_cell(first_value(fields, TEMPLATE_TYPE_KEYS))
        status = unwrap_feishu_cell(fields.get("生成状态") or fields.get("status"))
        if not emp_id and not depart_date and not return_date and not passport and not template_type:
            continue
        records.append({
            "record_id": item.get("record_id") or item.get("id") or "",
            "工号": emp_id.zfill(3) if emp_id else "",
            "护照号": passport if passport else "N/A",
            "出发时间": depart_date,
            "返程时间": return_date,
            "模板类型": template_type,
            "生成状态": status,
        })
    return records


def query_employee(emp_id):
    """从 SQLite 知识库查询人员信息"""
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT name, department, position, gender, passport FROM employees WHERE emp_id = ?", (emp_id,))
    result = cursor.fetchone()
    conn.close()
    if result:
        return {"name": result[0], "department": result[1], "position": result[2], "gender": result[3], "passport": result[4] if result[4] else ""}
    return None


def normalize_passport(value):
    """Return a comparable passport value. Empty / N/A / NaN are treated as missing."""
    if value is None:
        return ""
    text = str(value).strip()
    if not text or text.lower() == "nan" or text.upper() == "N/A":
        return ""
    return text.replace(" ", "").upper()


def find_employee_by_passport(passport, exclude_emp_id=None):
    passport = normalize_passport(passport)
    if not passport:
        return None
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    if exclude_emp_id:
        cursor.execute(
            "SELECT emp_id, name FROM employees WHERE UPPER(REPLACE(passport, ' ', '')) = ? AND emp_id != ?",
            (passport, exclude_emp_id),
        )
    else:
        cursor.execute(
            "SELECT emp_id, name FROM employees WHERE UPPER(REPLACE(passport, ' ', '')) = ?",
            (passport,),
        )
    result = cursor.fetchone()
    conn.close()
    if result:
        return {"emp_id": result[0], "name": result[1]}
    return None


def update_employee_passport(emp_id, passport):
    passport = normalize_passport(passport)
    if not passport:
        return False
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("UPDATE employees SET passport = ? WHERE emp_id = ?", (passport, emp_id))
    conn.commit()
    changed = cursor.rowcount > 0
    conn.close()
    return changed


def canonical_value(value):
    if pd.isna(value):
        return ""
    if isinstance(value, (pd.Timestamp,)):
        return value.strftime("%Y-%m-%d")
    text = str(value).strip()
    if " " in text and re_match_date(text):
        return text.split()[0]
    return text


def re_match_date(text):
    import re
    return bool(re.match(r"^\d{4}-\d{2}-\d{2}", text))


def tracking_id_for_row(row):
    record_id = str(row.get("record_id", "")).strip()
    if record_id:
        return f"record:{record_id}"
    return legacy_tracking_id_for_row(row)


def legacy_tracking_id_for_row(row):
    return f"legacy:{row['工号']}_{canonical_value(row['出发时间'])}"


def fingerprint_for_row(row):
    fields = {
        "emp_id": canonical_value(row.get("工号")),
        "name": canonical_value(row.get("name")),
        "department": canonical_value(row.get("department")),
        "position": canonical_value(row.get("position")),
        "gender": canonical_value(row.get("gender")),
        "passport": normalize_passport(row.get("passport")),
        "template_type": canonical_value(row.get("模板类型")),
        "depart_date": canonical_value(row.get("出发时间")),
        "return_date": canonical_value(row.get("返程时间")),
    }
    payload = json.dumps(fields, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def load_tracking():
    if not os.path.exists(TRACKING_FILE):
        return {}
    try:
        with open(TRACKING_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception:
        return {}
    if isinstance(data, dict) and "records" in data:
        return data.get("records", {})
    if isinstance(data, dict):
        return data
    if isinstance(data, list):
        # Backward compatible: old generated.json was a list of generated keys.
        return {f"legacy:{item}": {"fingerprint": None, "legacy": True} for item in data}
    return {}


def save_tracking(records):
    payload = {
        "version": 2,
        "records": records,
    }
    with open(TRACKING_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def query_employees_by_ids(emp_ids):
    """批量查询人员信息"""
    if not emp_ids:
        return {}
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    placeholders = ','.join(['?' for _ in emp_ids])
    cursor.execute(f"SELECT emp_id, name, department, position, gender FROM employees WHERE emp_id IN ({placeholders})", emp_ids)
    results = cursor.fetchall()
    conn.close()
    return {row[0]: {"name": row[1], "department": row[2], "position": row[3], "gender": row[4]} for row in results}


# === 工具函数 ===
def set_run_font(run, font_name=FONT_NAME, size_pt=FONT_SIZE_PT):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_paragraph_font(para, font_name=FONT_NAME, size_pt=FONT_SIZE_PT):
    for run in para.runs:
        set_run_font(run, font_name, size_pt)


def to_pinyin_upper(name_zh):
    name = str(name_zh).strip()
    parts = []
    i = 0
    while i < len(name):
        if i + 1 < len(name) and name[i:i+2] in PINYIN_DATA:
            parts.append(PINYIN_DATA[name[i:i+2]])
            i += 2
        elif name[i] in PINYIN_DATA:
            parts.append(PINYIN_DATA[name[i]])
            i += 1
        elif '一' <= name[i] <= '鿿':
            parts.append(name[i])
            print(f"  [警告] 未收录拼音: {name[i]} (来自 '{name}')")
            i += 1
        else:
            parts.append(name[i])
            i += 1
    if len(parts) >= 2:
        return f"{parts[0].upper()} {''.join(parts[1:]).upper()}"
    return "".join(parts).upper()


def fmt_date(dt):
    return f"{dt.strftime('%B')} {dt.day}, {dt.year}"


def fmt_short_date(dt):
    return f"{dt.strftime('%B')} {dt.day}, {dt.year}"


def resolve_dept(dept_raw, warnings):
    dept = str(dept_raw).strip()
    if dept in DEPT_TEMPLATE_MAP:
        return DEPT_TEMPLATE_MAP[dept]
    for key, val in DEPT_TEMPLATE_MAP.items():
        if dept.startswith(key) or key.startswith(dept):
            return val
    warnings.append(dept)
    return None


def resolve_template_for_row(row, warnings):
    template_type = canonical_value(row.get("模板类型", ""))
    if template_type:
        resolved = resolve_dept(template_type, warnings)
        if not resolved:
            print(f"  [警告] 工号 {row.get('工号', '')} 的模板类型“{template_type}”无法匹配模板")
        return resolved
    return resolve_dept(row.get("department", ""), warnings)


def title_with_mr(name_zh, gender):
    py = to_pinyin_upper(name_zh)
    title = "Mr." if str(gender).strip() == "男" else "Ms."
    return f"{title} {py}"


def translate_pos(pos, emp_id=None, name=None):
    """翻译职位，无法翻译则打印警告并保留原文"""
    if not isinstance(pos, str) or not pos.strip():
        return "N/A"
    pos = pos.strip()
    if pos in POSITION_MAP:
        return POSITION_MAP[pos]
    # 按 key 长度降序匹配，避免短 key 吞掉长 key（如 "工程师" 误匹配 "资深基建工程师"）
    for cn, en in sorted(POSITION_MAP.items(), key=lambda x: -len(x[0])):
        if cn in pos:
            return en
    # 翻译失败：打印警告，便于定位缺失的映射条目
    identifier = f"工号 {emp_id} ({name})" if emp_id and name else "未知"
    print(f"  [警告] {identifier}: 职位 '{pos}' 未在 POSITION_MAP 中找到翻译，保留原文")
    return pos


# === 文档生成 ===
def fill_personnel(doc, people, is_reply=False):
    def _build_inv_blocks(ppl):
        blocks = []
        for p in ppl:
            name = to_pinyin_upper(p["name"])
            blocks.append([f"Name: {name}", f"Passport Number: {p['passport']}", f"Title: {p['position']}"])
        return blocks

    def _build_rep_lines(ppl):
        lines = []
        for p in ppl:
            name = title_with_mr(p["name"], p["gender"])
            lines.append(f"{name}, Position: {p['position']}, Passport No.: {p['passport']}")
        return lines

    if is_reply:
        return _fill_personnel_reply(doc, people, _build_rep_lines(people))
    else:
        return _fill_personnel_inv(doc, people, _build_inv_blocks(people))


def _fill_personnel_inv(doc, people, blocks):
    name_idxs, passport_idxs, pos_idxs = [], [], []
    for i, p in enumerate(doc.paragraphs):
        if "{{NAME}}" in p.text: name_idxs.append(i)
        if "{{PASSPORT}}" in p.text: passport_idxs.append(i)
        if "{{POSITION}}" in p.text: pos_idxs.append(i)

    slot_count = min(len(name_idxs), len(passport_idxs), len(pos_idxs))
    if slot_count == 0:
        return doc

    filled_count = min(slot_count, len(people))
    for s in range(filled_count):
        _replace_para_text(doc.paragraphs[name_idxs[s]], blocks[s][0])
        _replace_para_text(doc.paragraphs[passport_idxs[s]], blocks[s][1])
        _replace_para_text(doc.paragraphs[pos_idxs[s]], blocks[s][2])

    if filled_count < slot_count:
        to_delete = []
        for s in range(filled_count, slot_count):
            to_delete.extend([name_idxs[s], passport_idxs[s], pos_idxs[s]])
            next_idx = pos_idxs[s] + 1
            while next_idx < len(doc.paragraphs) and not doc.paragraphs[next_idx].text.strip():
                to_delete.append(next_idx)
                next_idx += 1
        for idx in sorted(to_delete, reverse=True):
            doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    if len(people) > slot_count:
        from copy import deepcopy
        # 使用第一个模板槽位作为样式参考
        ref_para = doc.paragraphs[name_idxs[0]]
        last_idx = max(name_idxs[filled_count - 1], passport_idxs[filled_count - 1], pos_idxs[filled_count - 1])
        last_elem = doc.paragraphs[last_idx]._element

        for extra_p in people[slot_count:]:
            for line in [f"Name: {to_pinyin_upper(extra_p['name'])}", f"Passport Number: {extra_p['passport']}", f"Title: {extra_p['position']}"]:
                # 深拷贝模板段落的XML结构
                new_elem = deepcopy(ref_para._element)
                last_elem.addnext(new_elem)
                last_elem = new_elem
                new_para = type(doc.paragraphs[0])(new_elem, doc.paragraphs[0]._parent)
                # 替换文本保留格式
                from docx.oxml.ns import qn as qn2
                for run in new_para.runs:
                    run.text = ''
                if new_para.runs:
                    new_para.runs[0].text = line
                else:
                    new_para.text = line
                set_paragraph_font(new_para)

    # 处理 {{INVITE_PERSONNEL_LIST}} 占位符（人事组等模板使用）
    # 只列出超出单独槽位的人员，无人超出则删除占位符段落
    for i, p in enumerate(doc.paragraphs):
        if "{{INVITE_PERSONNEL_LIST}}" in p.text:
            remaining = people[filled_count:]  # 超出单独槽位的人员
            if remaining:
                list_lines = []
                for person in remaining:
                    name = to_pinyin_upper(person["name"])
                    list_lines.append(f"Name: {name}")
                    list_lines.append(f"Passport Number: {person['passport']}")
                    list_lines.append(f"Title: {person['position']}")
                _replace_para_text(p, "\n".join(list_lines))
            else:
                p._element.getparent().remove(p._element)
            break

    return doc


def _fill_personnel_reply(doc, people, lines):
    name_idxs = [i for i, p in enumerate(doc.paragraphs) if "{{NAME}}" in p.text]
    if not name_idxs:
        return doc

    body_idxs = [i for i in name_idxs if "CC:" not in doc.paragraphs[i].text]
    cc_idxs = [i for i in name_idxs if "CC:" in doc.paragraphs[i].text]

    if cc_idxs:
        cc_names = ", ".join(title_with_mr(p["name"], p["gender"]) for p in people)
        for cc_i in cc_idxs:
            _replace_para_text(doc.paragraphs[cc_i], f"CC: {cc_names}")

    filled_body = min(len(body_idxs), len(lines))
    for s in range(filled_body):
        _replace_para_text(doc.paragraphs[body_idxs[s]], lines[s])

    if filled_body < len(body_idxs):
        for idx in sorted(body_idxs[filled_body:], reverse=True):
            doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

    # 处理 {{INVITE_PERSONNEL_LIST}} 占位符
    has_personnel_list = False
    for i, p in enumerate(doc.paragraphs):
        if "{{INVITE_PERSONNEL_LIST}}" in p.text:
            has_personnel_list = True
            remaining = people[filled_body:]
            if remaining:
                extra_lines = []
                for person in remaining:
                    formatted = f"{title_with_mr(person['name'], person['gender'])}, Position: {person['position']}, Passport No.: {person['passport']}"
                    extra_lines.append(formatted)
                _replace_para_text(p, "\n".join(extra_lines))
            else:
                p._element.getparent().remove(p._element)
            break

    if not has_personnel_list and len(people) > len(body_idxs):
        from copy import deepcopy
        ref_para = doc.paragraphs[body_idxs[0]]
        last_elem = doc.paragraphs[body_idxs[-1]]._element
        for extra_line in lines[len(body_idxs):]:
            new_elem = deepcopy(ref_para._element)
            last_elem.addnext(new_elem)
            last_elem = new_elem
            new_para = type(doc.paragraphs[0])(new_elem, doc.paragraphs[0]._parent)
            for run in new_para.runs:
                run.text = ''
            if new_para.runs:
                new_para.runs[0].text = extra_line
            else:
                new_para.text = extra_line
            set_paragraph_font(new_para)
    return doc


def _replace_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.text = new_text
    set_paragraph_font(para)


def cleanup_date_artifacts(para):
    import re
    full_text = ''.join(r.text for r in para.runs)
    # 移除模板中日期占位符后面的硬编码年份（如 {{DATE}}, 2026 → 替换后变 "May 20, 2026, 2026"）
    # 匹配任意年份重复：2026, 2026 / 2027, 2027 等，保留第一个年份
    full_text = re.sub(r'(\d{4})[, ]+\1', r'\1', full_text)
    if para.runs:
        para.runs[0].text = full_text
        for run in para.runs[1:]:
            run.text = ''


def replace_non_personnel(doc, replacements):
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text and "{{NAME}}" not in key and "{{PASSPORT}}" not in key and "{{POSITION}}" not in key:
                found_in_runs = any(key in run.text for run in para.runs)
                if found_in_runs:
                    for run in para.runs:
                        if key in run.text:
                            run.text = run.text.replace(key, val)
                            set_run_font(run)
                else:
                    # 占位符跨多个 run 拆分（如 "{" + "{TRAVEL_DATE}}"），整段替换
                    _replace_para_text(para, para.text.replace(key, val))
        cleanup_date_artifacts(para)
    return doc


def generate_for_dept(dept_folder, dept_label, records, names_str=""):
    inv_path = os.path.join(TEMPLATE_DIR, dept_folder, "邀请函.docx")
    rep_path = os.path.join(TEMPLATE_DIR, dept_folder, "答复函.docx")
    if not os.path.exists(inv_path) or not os.path.exists(rep_path):
        print(f"  [跳过] 模板不存在: {dept_folder}")
        return

    travel_start = records["adj_depart"].min()
    travel_end = records["adj_return"].max()
    issue_date = travel_start - relativedelta(months=1)
    reply_date = issue_date + pd.Timedelta(days=2)
    if travel_start == travel_end:
        travel_period = fmt_short_date(travel_start)
    else:
        travel_period = f"{fmt_short_date(travel_start)} to {fmt_short_date(travel_end)}"

    people = [{"name": row["name"], "passport": row["passport"], "position": row["position"], "gender": row["gender"]} for _, row in records.iterrows()]

    year_str = str(travel_start.year)
    common_replacements = {
        "{{ISSUE_DATE}}": fmt_date(issue_date),
        "{{TRAVEL_DATE}}": travel_period,
        "{{TravelPeriod}}": f"{travel_period}, {year_str}",
    }
    rep_extra = {"{{REPLY_DATE}}": fmt_date(reply_date)}

    inv_filename = f"{dept_label}_邀请函_{names_str}.docx" if names_str else f"{dept_label}_邀请函.docx"
    rep_filename = f"{dept_label}_答复函_{names_str}.docx" if names_str else f"{dept_label}_答复函.docx"

    inv_doc = Document(inv_path)
    replace_non_personnel(inv_doc, common_replacements)
    fill_personnel(inv_doc, people, is_reply=False)
    inv_out = os.path.join(OUTPUT_DIR, inv_filename)
    inv_doc.save(inv_out)
    print(f"  [OK] 邀请函: {inv_out}")

    rep_doc = Document(rep_path)
    replace_non_personnel(rep_doc, {**common_replacements, **rep_extra})
    fill_personnel(rep_doc, people, is_reply=True)
    rep_out = os.path.join(OUTPUT_DIR, rep_filename)
    rep_doc.save(rep_out)
    print(f"  [OK] 答复函: {rep_out}")


# === 主流程 ===
def main():
    global AUTO_YES, DRY_RUN, FORCE_REGENERATE, ALLOW_MISSING_PASSPORT
    excel_path = None
    feishu_data_file = None

    # Parse arguments
    i = 1
    while i < len(sys.argv):
        if sys.argv[i] == '--yes':
            AUTO_YES = True
            i += 1
        elif sys.argv[i] == '--dry-run':
            DRY_RUN = True
            i += 1
        elif sys.argv[i] == '--force':
            FORCE_REGENERATE = True
            i += 1
        elif sys.argv[i] == '--allow-missing-passport':
            ALLOW_MISSING_PASSPORT = True
            i += 1
        elif sys.argv[i] == '--feishu' and i + 1 < len(sys.argv):
            feishu_data_file = sys.argv[i + 1]
            i += 2
        elif sys.argv[i] == '--excel' and i + 1 < len(sys.argv):
            excel_path = sys.argv[i + 1]
            i += 2
        elif not sys.argv[i].startswith('--'):
            excel_path = sys.argv[i]
            i += 1
        else:
            i += 1

    if not os.path.exists(DB_PATH):
        print(f"[错误] 知识库不存在: {DB_PATH}")
        print("请先运行 init_database.py 初始化知识库")
        sys.exit(1)

    # If Feishu data file is provided, use it directly
    if feishu_data_file:
        with open(feishu_data_file, 'r', encoding='utf-8') as f:
            raw_data = json.load(f)
        _process_feishu_data(raw_data)
        return

    # If --excel is specified, force Excel mode
    if excel_path:
        if not os.path.exists(excel_path):
            print(f"[错误] 文件不存在: {excel_path}")
            sys.exit(1)
        _process_excel(excel_path)
        return

    # Default: read from Feishu table directly
    if not FEISHU_BASE_TOKEN or not FEISHU_TABLE_ID:
        print("[提示] 请先设置环境变量:")
        print("  setx FEISHU_BASE_TOKEN \"YOUR_BASE_TOKEN\"")
        print("  setx FEISHU_TABLE_ID \"YOUR_TABLE_ID\"")
        print("  setx LARK_CLI_PATH \"C:\\path\\to\\lark-cli.cmd\"")
        print("  设置后重新打开终端即可生效")
        sys.exit(1)

    records = read_feishu_table(FEISHU_BASE_TOKEN, FEISHU_TABLE_ID)
    if records:
        raw_data = []
        for d in records:
            raw_data.append({
                'emp_id': d['工号'],
                'passport': d['护照号'] if d['护照号'] != 'N/A' else '',
                'depart_date': d['出发时间'],
                'return_date': d['返程时间'],
                'template_type': d.get('模板类型', ''),
                'record_id': d.get('record_id', ''),
                'generation_status': d.get('生成状态', '')
            })
        _process_feishu_data(raw_data)
    else:
        print("飞书表格无数据或读取失败")
        print("[错误] 为避免误生成旧数据，默认不会自动改读本地 Excel。")
        print("如需使用本地 Excel，请在命令中明确指定 --excel 文件路径。")
        sys.exit(1)


def _process_feishu_data(raw_data):
    """处理飞书数据"""
    print(f"飞书表格共 {len(raw_data)} 条记录")
    print("查询知识库...")

    records = []
    warnings = []
    missings = []  # 工号不存在的记录
    passport_updates = []
    passport_conflicts = []

    for d in raw_data:
        emp_id = d.get('emp_id', d.get('工号', ''))
        if not emp_id.startswith('00'):
            emp_id = emp_id.zfill(3)

        emp_info = query_employee(emp_id)
        if not emp_info:
            missings.append(f"工号 {emp_id} 不在知识库中，无法生成")
            print(f"  [警告] 工号 {emp_id} 不在知识库中，跳过")
            continue

        # 护照号优先级: 飞书表格 > 知识库；飞书补全后可自动写回空白知识库。
        feishu_passport = normalize_passport(d.get('passport', d.get('护照号', '')))
        kb_passport = normalize_passport(emp_info.get('passport', ''))

        if feishu_passport:
            other = find_employee_by_passport(feishu_passport, exclude_emp_id=emp_id)
            if other:
                passport_conflicts.append(
                    f"工号 {emp_id} ({emp_info['name']}) 的护照号 {feishu_passport} 已属于工号 {other['emp_id']} ({other['name']})，请人工核对"
                )
                update_feishu_status(d.get('record_id', ''), '护照冲突')
                continue
            elif kb_passport and feishu_passport != kb_passport:
                if update_employee_passport(emp_id, feishu_passport):
                    passport_updates.append(
                        f"工号 {emp_id} ({emp_info['name']}) 护照号已按飞书从 {kb_passport} 更新为 {feishu_passport}"
                    )
                    emp_info['passport'] = feishu_passport
            elif not kb_passport:
                if update_employee_passport(emp_id, feishu_passport):
                    passport_updates.append(f"工号 {emp_id} ({emp_info['name']}) 已同步护照号 {feishu_passport} 到知识库")
                    emp_info['passport'] = feishu_passport
            passport = feishu_passport
        elif kb_passport:
            passport = kb_passport
        else:
            update_feishu_status(d.get('record_id', ''), '待补护照')
            warnings.append(f"工号 {emp_id} ({emp_info['name']}) 护照号为空，请补充")
            if not ALLOW_MISSING_PASSPORT:
                continue
            passport = 'N/A'

        records.append({
            '工号': emp_id,
            'name': emp_info['name'],
            'department': emp_info['department'],
            '模板类型': d.get('template_type') or d.get('模板类型', ''),
            'position': translate_pos(emp_info['position'], emp_id, emp_info['name']),
            'gender': emp_info['gender'],
            'passport': passport,
            '出发时间': d.get('depart_date', d.get('出发时间', '')),
            '返程时间': d.get('return_date', d.get('返程时间', '')),
            'record_id': d.get('record_id', ''),
            '生成状态': d.get('generation_status', d.get('生成状态', ''))
        })

    # 输出警告 + 飞书通知
    if warnings:
        print("\n=== 护照号缺失警告 ===")
        for w in warnings:
            print(f"  [警告] {w}")
        notify_feishu("护照号缺失提醒", "\n".join(warnings))

    if passport_updates:
        print("\n=== 护照号已同步到知识库 ===")
        for item in passport_updates:
            print(f"  [OK] {item}")
        notify_feishu("护照号已同步到知识库", "\n".join(passport_updates))

    if passport_conflicts:
        print("\n=== 护照号冲突/陌生提醒 ===")
        for item in passport_conflicts:
            print(f"  [警告] {item}")
        notify_feishu("护照号冲突/陌生提醒", "\n".join(passport_conflicts))

    if missings:
        print("\n=== 工号缺失警告 ===")
        for m in missings:
            print(f"  [警告] {m}")
        notify_feishu("工号缺失提醒 - 请补充知识库", "\n".join(missings))

    if not records:
        print("[错误] 无有效数据")
        sys.exit(1)

    df = pd.DataFrame(records)
    df['depart_date'] = pd.to_datetime(df['出发时间'])
    df['return_date'] = pd.to_datetime(df['返程时间'])
    _process_dataframe(df)

def _process_dataframe(df):
    """处理 DataFrame 数据，生成邀请函"""
    warnings = []
    if '模板类型' not in df.columns:
        df['模板类型'] = ''
    df['__dept'] = df.apply(lambda row: resolve_template_for_row(row, warnings), axis=1)
    df['adj_depart'] = pd.to_datetime(df['depart_date']) - pd.Timedelta(days=1)
    df['adj_return'] = pd.to_datetime(df['return_date']) + pd.Timedelta(days=1)

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    tracking_records = load_tracking()

    new_depts = {}
    tracking_keys = []  # 记录本次生成的追踪 key 和 record_id
    migrated_tracking = []
    for _, row in df.iterrows():
        tracking_key = tracking_id_for_row(row)
        legacy_key = legacy_tracking_id_for_row(row)
        fingerprint = fingerprint_for_row(row)
        previous = tracking_records.get(tracking_key) or tracking_records.get(legacy_key)
        generation_status = canonical_value(row.get("生成状态", ""))
        if previous and previous.get("legacy") and not FORCE_REGENERATE:
            tracking_records[tracking_key] = {
                "fingerprint": fingerprint,
                "record_id": row.get("record_id", ""),
                "emp_id": row["工号"],
                "generated_at": pd.Timestamp.now().isoformat(),
                "migrated_from_legacy": True,
            }
            migrated_tracking.append(tracking_key)
            needs_generation = False
        elif previous is None and generation_status == "已生成" and not FORCE_REGENERATE:
            tracking_records[tracking_key] = {
                "fingerprint": fingerprint,
                "record_id": row.get("record_id", ""),
                "emp_id": row["工号"],
                "generated_at": pd.Timestamp.now().isoformat(),
                "migrated_from_feishu_status": True,
            }
            migrated_tracking.append(tracking_key)
            needs_generation = False
        else:
            needs_generation = FORCE_REGENERATE or previous is None or previous.get("fingerprint") != fingerprint
        if needs_generation:
            dept = row['__dept']
            if not dept:
                update_feishu_status(row.get('record_id', ''), '未知模板/部门')
                continue
            if dept not in new_depts:
                new_depts[dept] = []
            new_depts[dept].append(row)
            tracking_keys.append({
                'key': tracking_key,
                'fingerprint': fingerprint,
                'record_id': row.get('record_id', ''),
                'emp_id': row['工号']
            })

    if migrated_tracking:
        save_tracking(tracking_records)
        print(f"\n已同步 {len(migrated_tracking)} 条飞书已生成记录到本地追踪文件。")

    if not new_depts:
        print("\n没有新的人员需要生成，所有文件已是最新。")
        return

    total_new = sum(len(persons) for persons in new_depts.values())
    print(f"\n检测到 {total_new} 个新增人员:")
    for dept, persons in new_depts.items():
        names = [row['name'] for row in persons]
        print(f"  [{dept}] {len(persons)} 人: {', '.join(names)}")

    print()

    print("=== 生成预览 ===")
    for dept, persons in new_depts.items():
        dept_df = pd.DataFrame(persons)
        travel_start = dept_df["adj_depart"].min()
        travel_end = dept_df["adj_return"].max()
        issue_date = travel_start - relativedelta(months=1)
        reply_date = issue_date + pd.Timedelta(days=2)
        print(
            f"  [{dept}] {len(persons)} 人 | ISSUE_DATE={fmt_date(issue_date)} | "
            f"REPLY_DATE={fmt_date(reply_date)} | TRAVEL={fmt_date(travel_start)} to {fmt_date(travel_end)}"
        )
        for _, person in dept_df.iterrows():
            template_note = canonical_value(person.get('模板类型', '')) or f"按知识库部门: {person['department']}"
            print(f"    - {person['工号']} {person['name']} | 模板 {template_note} -> {dept} | 护照 {person['passport']} | 职位 {person['position']}")

    if DRY_RUN:
        print("\n[dry-run] 仅预览，不生成 Word 文件。")
        print("[dry-run] 如飞书已有“已生成”记录，可能会同步到 generated.json 作为后续变更判断依据。")
        return

    if not AUTO_YES:
        answer = input("\n确认生成以上邀请函/答复函？输入 YES 继续: ").strip()
        if answer != "YES":
            print("已取消生成。")
            return

    for dept_folder, persons in new_depts.items():
        dept_df = pd.DataFrame(persons)
        print(f"--- {dept_folder} ({len(persons)} 人) ---")
        person_names = [to_pinyin_upper(row['name']) for _, row in dept_df.iterrows()]
        names_str = '_'.join(person_names)
        generate_for_dept(dept_folder, dept_folder, dept_df, names_str)

    # 更新本地追踪文件：按 record_id/legacy key 保存关键字段指纹。
    generated_at = pd.Timestamp.now().isoformat()
    for k in tracking_keys:
        tracking_records[k['key']] = {
            "fingerprint": k["fingerprint"],
            "record_id": k.get("record_id", ""),
            "emp_id": k.get("emp_id", ""),
            "generated_at": generated_at,
        }
    save_tracking(tracking_records)

    # 写回飞书表格标记"已生成"
    print("\n更新飞书表格生成状态...")
    for k in tracking_keys:
        update_feishu_status(k['record_id'], '已生成')

    if warnings:
        print(f"\n[警告] 以下模板类型或部门无匹配模板，已跳过: {', '.join(set(warnings))}")

    print(f"\n完成! 文件输出至: {OUTPUT_DIR}")
    print(f"已生成 {len(tracking_keys)} 人，已追踪 {len(tracking_records)} 条记录")


def _process_excel(excel_path):
    """从本地 Excel 处理数据"""
    df = pd.read_excel(excel_path, engine="openpyxl")
    df.columns = [str(c).replace("\n", "").strip() for c in df.columns]
    print(f"读取 Excel: {excel_path}")

    emp_ids = []
    for _, row in df.iterrows():
        emp_id = row['工号']
        if isinstance(emp_id, (int, float)):
            emp_ids.append(str(int(emp_id)).zfill(3))
        else:
            emp_ids.append(str(emp_id).strip().zfill(3))
    employees = query_employees_by_ids(emp_ids)

    merged_data = []
    for _, row in df.iterrows():
        emp_id_raw = row['工号']
        emp_id = str(int(emp_id_raw)).zfill(3) if isinstance(emp_id_raw, (int, float)) else str(emp_id_raw).strip().zfill(3)
        emp_info = employees.get(emp_id)
        if not emp_info:
            print(f"  [警告] 工号 {emp_id} 在知识库中不存在，跳过")
            continue
        merged_data.append({
            "emp_id": emp_id, "工号": emp_id,
            "name": emp_info["name"], "department": emp_info["department"],
            "模板类型": str(row['模板类型']).strip() if pd.notna(row.get('模板类型')) else "",
            "position": translate_pos(emp_info["position"], emp_id, emp_info["name"]), "gender": emp_info["gender"],
            "passport": str(row['护照号']).strip() if pd.notna(row.get('护照号')) else "N/A",
            "depart_date": row['出发时间'], "return_date": row['返程时间'],
            "出发时间": row['出发时间'], "返程时间": row['返程时间']
        })

    if not merged_data:
        print("[错误] 无有效数据")
        sys.exit(1)

    df_merged = pd.DataFrame(merged_data)
    _process_dataframe(df_merged)


if __name__ == "__main__":
    main()
